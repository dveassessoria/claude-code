"""
docs-base-cliente — helpers.py
Funções reutilizáveis de formatação para geração de documentos .docx de clientes DVE.

Uso no script do cliente:
    import sys
    sys.path.insert(0, '/Users/macbookairm4/Documents/DVE Assessoria/Claude Code/.claude/skills/docs-base-cliente')
    from helpers import *
    # Depois defina as cores do cliente e chame as funções.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = 'Montserrat'

# Cores padrão (sobrescrever no script do cliente)
COR_PRIMARIA   = RGBColor(0x1A, 0x1A, 0x1A)
COR_SECUNDARIA = RGBColor(0x44, 0x44, 0x44)
COR_TEXTO      = RGBColor(0x1E, 0x1E, 0x1E)
COR_CINZA_MED  = RGBColor(0x66, 0x66, 0x66)
COR_CINZA_LEVE = RGBColor(0xF2, 0xF2, 0xF2)
COR_BRANCO     = RGBColor(0xFF, 0xFF, 0xFF)
COR_VERDE      = RGBColor(0x1E, 0x7E, 0x34)
COR_AMARELO    = RGBColor(0xB8, 0x86, 0x00)
COR_VERMELHO   = RGBColor(0xC8, 0x10, 0x10)


# ── Células ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
    tcPr.append(shd)


def set_cell_border(cell, **kw):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        if edge in kw:
            tag.set(qn('w:val'), kw[edge].get('val', 'single'))
            tag.set(qn('w:sz'), str(kw[edge].get('sz', 4)))
            tag.set(qn('w:color'), kw[edge].get('color', 'auto'))
        else:
            tag.set(qn('w:val'), 'nil')
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=100, right=140, bottom=100, left=140):
    """Padding interno da célula (twips; 1pt ≈ 20twips)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('right', right), ('bottom', bottom), ('left', left)]:
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:w'), str(val))
        tag.set(qn('w:type'), 'dxa')
        tcMar.append(tag)
    tcPr.append(tcMar)


# ── Texto ─────────────────────────────────────────────────────────────────────

def run(para, text, bold=False, size=10.5, color=None, italic=False):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.color.rgb = color or COR_TEXTO
    return r


def heading1(doc, text, color=None):
    """Título de seção principal — usa COR_PRIMARIA por padrão."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    run(p, text, bold=True, size=16, color=color or COR_PRIMARIA)
    return p


def heading2(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run(p, text, bold=True, size=12, color=color or COR_SECUNDARIA)
    return p


def heading3(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run(p, text, bold=True, size=11, color=color or COR_CINZA_MED)
    return p


def body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run(p, text, size=10.5, italic=italic)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run(p, bold_prefix, bold=True, size=10.5)
        run(p, text, size=10.5)
    else:
        run(p, text, size=10.5)
    return p


def divider(doc, color=None):
    """Linha divisória horizontal na cor primária."""
    hex_color = '{:02X}{:02X}{:02X}'.format(*(color or COR_PRIMARIA))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ── Tabelas ───────────────────────────────────────────────────────────────────

def table_2col(doc, rows, header=None, header_color=None):
    """
    Tabela de 2 colunas com linhas alternadas.
    rows: lista de tuplas (col1, col2) ou (col1, (bold_prefix, rest))
    """
    n_rows = len(rows) + (1 if header else 0)
    t = doc.add_table(rows=n_rows, cols=2)
    t.style = 'Table Grid'
    hc = header_color or COR_PRIMARIA
    hc_hex = '{:02X}{:02X}{:02X}'.format(*hc)
    hdr_brd = {'val': 'single', 'sz': 10, 'color': hc_hex}
    bdy_brd = {'val': 'single', 'sz': 6,  'color': hc_hex}
    start = 0
    if header:
        for i, h in enumerate(header):
            cell = t.cell(0, i)
            set_cell_bg(cell, hc)
            set_cell_margins(cell)
            set_cell_border(cell, top=hdr_brd, bottom=hdr_brd, left=hdr_brd, right=hdr_brd)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run(p, h, bold=True, size=10, color=COR_BRANCO)
        start = 1
    for ri, row in enumerate(rows):
        bg = COR_CINZA_LEVE if ri % 2 == 0 else COR_BRANCO
        for ci, val in enumerate(row):
            cell = t.cell(ri + start, ci)
            set_cell_bg(cell, bg)
            set_cell_margins(cell)
            set_cell_border(cell, top=bdy_brd, bottom=bdy_brd, left=bdy_brd, right=bdy_brd)
            p = cell.paragraphs[0]
            if isinstance(val, tuple):
                run(p, val[0], bold=True, size=10)
                run(p, val[1], size=10)
            else:
                run(p, val, size=10)
    return t


def table_3col(doc, rows, header=None, header_color=None):
    """Tabela de 3 colunas com linhas alternadas."""
    n_rows = len(rows) + (1 if header else 0)
    t = doc.add_table(rows=n_rows, cols=3)
    t.style = 'Table Grid'
    hc = header_color or COR_PRIMARIA
    hc_hex = '{:02X}{:02X}{:02X}'.format(*hc)
    hdr_brd = {'val': 'single', 'sz': 10, 'color': hc_hex}
    bdy_brd = {'val': 'single', 'sz': 6,  'color': hc_hex}
    start = 0
    if header:
        for i, h in enumerate(header):
            cell = t.cell(0, i)
            set_cell_bg(cell, hc)
            set_cell_margins(cell)
            set_cell_border(cell, top=hdr_brd, bottom=hdr_brd, left=hdr_brd, right=hdr_brd)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run(p, h, bold=True, size=10, color=COR_BRANCO)
        start = 1
    for ri, row in enumerate(rows):
        bg = COR_CINZA_LEVE if ri % 2 == 0 else COR_BRANCO
        for ci, val in enumerate(row):
            cell = t.cell(ri + start, ci)
            set_cell_bg(cell, bg)
            set_cell_margins(cell)
            set_cell_border(cell, top=bdy_brd, bottom=bdy_brd, left=bdy_brd, right=bdy_brd)
            run(cell.paragraphs[0], str(val), bold=(ci == 0), size=10)
    return t


def semaforo_table(doc, rows, header_color=None):
    """
    Tabela de diagnóstico com semáforo.
    rows: lista de (elemento, status, cor_status, prioridade)
    """
    hc = header_color or COR_PRIMARIA
    hc_hex = '{:02X}{:02X}{:02X}'.format(*hc)
    hdr_brd = {'val': 'single', 'sz': 10, 'color': hc_hex}
    bdy_brd = {'val': 'single', 'sz': 6,  'color': hc_hex}
    t = doc.add_table(rows=len(rows) + 1, cols=3)
    t.style = 'Table Grid'
    for i, h in enumerate(['Elemento', 'Status atual', 'Prioridade']):
        cell = t.cell(0, i)
        set_cell_bg(cell, hc)
        set_cell_margins(cell)
        set_cell_border(cell, top=hdr_brd, bottom=hdr_brd, left=hdr_brd, right=hdr_brd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, h, bold=True, size=10, color=COR_BRANCO)
    for ri, (elem, status, cor, prio) in enumerate(rows):
        bg = COR_CINZA_LEVE if ri % 2 == 0 else COR_BRANCO
        for ci, (val, kw) in enumerate([
            (elem,   {'bold': True}),
            (status, {'color': cor}),
            (prio,   {}),
        ]):
            cell = t.cell(ri + 1, ci)
            set_cell_bg(cell, bg)
            set_cell_margins(cell)
            set_cell_border(cell, top=bdy_brd, bottom=bdy_brd, left=bdy_brd, right=bdy_brd)
            run(cell.paragraphs[0], val, size=10, **kw)
    return t


# ── Capa ──────────────────────────────────────────────────────────────────────

def cover_page(doc, company_name, subtitle, sections_summary, date_str):
    """
    Gera página de capa padrão DVE.
    sections_summary: string com os títulos das seções separados por '  •  '
    date_str: ex. 'Junho de 2026'
    """
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, company_name.upper(), bold=True, size=22, color=COR_PRIMARIA)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, subtitle, size=12, color=COR_CINZA_MED, italic=True)

    doc.add_paragraph()
    divider(doc)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, 'BASE DO PROJETO', bold=True, size=18, color=COR_SECUNDARIA)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, sections_summary, size=10, color=COR_CINZA_MED, italic=True)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, f'Elaborado por DVE Assessoria — {date_str}', size=10, color=COR_CINZA_MED, italic=True)

    doc.add_page_break()


# ── Documento base ────────────────────────────────────────────────────────────

def new_doc():
    """Cria documento com margens padrão DVE."""
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(3)
    return doc
